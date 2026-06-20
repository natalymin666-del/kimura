#!/usr/bin/env python3
"""
Kimura Scope Agreement Generator
==================================

Generates a basic Scope of Work / Testing Authorization document for an
LLM security engagement. This is a STARTING POINT, not a finished legal
contract — have it reviewed by an actual lawyer (ideally one familiar
with NL/EU IT services contracts) before using it with a real paying
client. This script exists to save you from writing the boilerplate by
hand each time, not to replace legal review.

USAGE:
    python3 kimura_scope_agreement.py \\
        --client "Acme Corp" \\
        --client-contact "Jane Doe, CTO" \\
        --systems "Production customer support chatbot (chat.acme.com)" \\
        --start-date 2026-07-01 \\
        --end-date 2026-07-14 \\
        --fee "2500 EUR" \\
        --output acme_scope_agreement.docx
"""

import argparse
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return h


def build_agreement(args):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("LLM Security Testing — Scope of Work & Authorization", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Prepared {datetime.now().strftime('%Y-%m-%d')}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ---- Parties ----
    add_heading(doc, "1. Parties", level=1)
    p = doc.add_paragraph()
    p.add_run("Service Provider: ").bold = True
    p.add_run(f"{args.provider_name}\n")
    p.add_run("Client: ").bold = True
    p.add_run(f"{args.client}")
    if args.client_contact:
        p.add_run(f" (Contact: {args.client_contact})")

    # ---- Scope ----
    add_heading(doc, "2. Scope of Testing", level=1)
    doc.add_paragraph(
        "The Client explicitly authorizes the Service Provider to perform security "
        "testing limited to the following system(s), and no others:"
    )
    p = doc.add_paragraph(args.systems, style="List Bullet")

    doc.add_paragraph(
        "Testing is limited to the following techniques unless otherwise agreed in "
        "writing in advance:"
    )
    for item in [
        "Prompt injection testing (direct and indirect)",
        "System prompt / configuration extraction attempts",
        "Jailbreak and safety-guardrail bypass testing",
        "Testing of LLM-specific behaviors as defined in the OWASP Top 10 for LLM Applications",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "OUT OF SCOPE (explicitly excluded unless separately authorized in writing): "
        "any testing against systems, accounts, or data not listed above; any attempt "
        "to access, exfiltrate, modify, or delete real user/customer data; any denial-"
        "of-service or load-based testing; any testing against third-party "
        "infrastructure the Client does not own or control; social engineering of "
        "Client personnel."
    )

    # ---- Testing window ----
    add_heading(doc, "3. Testing Window", level=1)
    p = doc.add_paragraph()
    p.add_run("Start date: ").bold = True
    p.add_run(f"{args.start_date}\n")
    p.add_run("End date: ").bold = True
    p.add_run(f"{args.end_date}\n")
    doc.add_paragraph(
        "Testing authorization under this agreement is valid only within this window. "
        "Any testing outside this window requires a new written authorization."
    )

    # ---- Methodology & reporting ----
    add_heading(doc, "4. Methodology & Deliverables", level=1)
    doc.add_paragraph(
        "Testing will use independent repeated trials per identified attack vector "
        "(typically n=30) with statistical reproducibility reporting (Wilson 95% "
        "confidence intervals), rather than single-occurrence demonstrations. Where a "
        "vulnerability is identified, the Service Provider will, where feasible, also "
        "test and report on at least one candidate mitigation."
    )
    doc.add_paragraph(
        "Deliverable: a written report including an executive summary, severity "
        "assessment, reproducibility data, example findings (redacted of any real "
        "Client data if encountered unexpectedly), and recommendations."
    )

    # ---- Confidentiality ----
    add_heading(doc, "5. Confidentiality", level=1)
    doc.add_paragraph(
        "The Service Provider agrees to keep confidential any non-public information "
        "about the Client's systems, including but not limited to system prompts, "
        "architecture details, vulnerabilities discovered, and any data encountered "
        "during testing. This obligation survives termination of this agreement. "
        "Findings will not be published, disclosed, or referenced publicly (including "
        "in portfolios, blog posts, or social media) without the Client's prior written "
        "consent, which may be limited to an anonymized or generalized description of "
        "the work performed."
    )

    # ---- Data handling ----
    add_heading(doc, "6. Data Handling")
    doc.add_paragraph(
        "If real Client or end-user data is unexpectedly encountered during testing, "
        "the Service Provider will immediately stop, document the minimum necessary "
        "detail to report the issue, securely delete any such data from local systems, "
        "and notify the Client within 24 hours."
    )

    # ---- Liability ----
    add_heading(doc, "7. Limitation of Liability", level=1)
    doc.add_paragraph(
        "Testing carries an inherent risk of unexpected system behavior. The Client "
        "acknowledges this risk and agrees that the Service Provider's liability for "
        "any damages arising from this engagement is limited to the fees paid under "
        "this agreement, except in cases of gross negligence or willful misconduct. "
        "This clause should be reviewed by the Client's own legal counsel and adjusted "
        "as needed for the applicable jurisdiction."
    )

    # ---- Fees ----
    add_heading(doc, "8. Fees & Payment", level=1)
    p = doc.add_paragraph()
    p.add_run("Fee: ").bold = True
    p.add_run(f"{args.fee}\n")
    doc.add_paragraph(
        "Payment terms, invoicing schedule, and currency to be confirmed separately in "
        "writing prior to commencement of testing."
    )

    # ---- Signatures ----
    add_heading(doc, "9. Authorization", level=1)
    doc.add_paragraph(
        "By signing below, the Client confirms it has the legal authority to authorize "
        "security testing of the system(s) listed in Section 2, and grants that "
        "authorization to the Service Provider for the testing window specified in "
        "Section 3."
    )

    doc.add_paragraph()
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.cell(0, 0).text = "Service Provider signature:"
    sig_table.cell(0, 1).text = "Client signature:"
    sig_table.cell(1, 0).text = "_______________________"
    sig_table.cell(1, 1).text = "_______________________"
    sig_table.cell(2, 0).text = "Date:"
    sig_table.cell(2, 1).text = "Date:"

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(
        "This document is a starting template, not finished legal advice. Have it "
        "reviewed by a qualified lawyer before use with a real client."
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    fr.italic = True

    doc.save(args.output)
    return args.output


def main():
    parser = argparse.ArgumentParser(description="Generate a Scope of Work / Testing Authorization document.")
    parser.add_argument("--provider-name", default="Kimura — LLM Security Research", help="Your business/brand name")
    parser.add_argument("--client", required=True, help="Client company name")
    parser.add_argument("--client-contact", default="", help="Client contact person and title")
    parser.add_argument("--systems", required=True, help="Description of in-scope system(s)")
    parser.add_argument("--start-date", required=True, help="Testing window start date")
    parser.add_argument("--end-date", required=True, help="Testing window end date")
    parser.add_argument("--fee", default="To be agreed", help="Fee amount and currency")
    parser.add_argument("--output", default="scope_agreement.docx", help="Output .docx path")
    args = parser.parse_args()

    path = build_agreement(args)
    print(f"Scope agreement saved to {path}")
    print("REMINDER: have this reviewed by a lawyer before sending to a real client.")


if __name__ == "__main__":
    main()
